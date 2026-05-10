from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "echo-protocol"
OUT_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT_DIR / "Echo-Protocol-Asset-Brief.docx"
PDF_PATH = OUT_DIR / "Echo-Protocol-Asset-Brief.pdf"

TITLE = "Echo Protocol\nAsset Brief / Cutscenes, Faux-Real Material, Audio"
SUBTITLE = (
    "Ziel: Echo Protocol soll sich wie ein realer Fall mit abgefilmten Beweisen, "
    "Scans, Abhoer-Fragmenten und dokumentarischen Fundstuecken anfuehlen - "
    "nicht wie plain UI-Text."
)
PRIORITY_TEXT = (
    "Prioritaet fuer den naechsten echten Realismus-Sprung:\n"
    "1. Standbilder / Scans fuer Cutscenes\n"
    "2. Kleine Audiofragmente oder Voiceover-Transcripts\n"
    "3. Faux-UI / Archiv- und Dossiergrafiken\n"
    "4. Optional spaeter: kurze Loop-Video-Schnipsel"
)

DROP_FOLDER_TEXT = (
    "Drop-in-Ordner fuer spaetere Assets:\n"
    "public/assets/echo-protocol/cutscenes/slots/\n"
    "Die Cutscenes pruefen dort automatisch feste Dateinamen pro Shot."
)

NEEDS = [
    "3-5 Hauptbilder im dokumentarischen Stil: Flur / CCTV, Akte / Schreibtisch, Mira im Rotlicht, Tower/Technik, Exit Shot.",
    "1-3 Charakterbilder oder Silhouetten: Mira, optional Jonas, optional Elias nur angedeutet.",
    "Optional kurze Voice-Snippets oder zumindest finaler Text fuer Mira, Dispatch, Jonas und Elias.",
    "Falls du echte Fotos oder Moodboards hast: auch Handyfotos, Screenshots, Midjourney-Stills oder AI-Bilder sind als Startpunkt okay.",
    "Wenn du gar keine Bilder hast, reichen mir zuerst Referenzen pro Shot mit Stimmung, Farbe, Perspektive und Kleidungs-/Raumhinweisen.",
]

SHOT_HEADERS = ["Shot", "Einsatz", "Was zu sehen ist", "Script / Gefuehl", "Dateiwunsch"]
SHOT_ROWS = [
    [
        "CCTV / Nordfluegel",
        "Spielintro",
        "Leerer Flur, Regenlicht, Polizeituer, statische Kamera",
        "Soll wie echte Ueberwachung wirken; kalt, beobachtend, minimal menschlich",
        "PNG/JPG, 1920x1080, 16:9",
    ],
    [
        "Desk / Dossier Scan",
        "Office + Akte",
        "Akte, Kaffeering, Polaroids, Haftnotizen, Stempel",
        "Nicht clean. Eher benutzt, als waere schon jemand mehrfach durch diese Nacht gegangen",
        "PNG/JPG hochaufgeloest, gern overhead",
    ],
    [
        "Signal Transcript",
        "Signal Trace",
        "Abhoer-/Band-Protokoll oder UI-artige Funkspur",
        "Rauschen, Fragment, etwas Offizielles, das leicht kippt",
        "Grafik oder Textgrundlage",
    ],
    [
        "Mira / Red Room Still",
        "Red Room Reveal",
        "Mira im roten Notlicht, halb real, halb falsch",
        "Mehr verbotener Beweis als Fantasy-Character-Art",
        "PNG/JPG, 16:9 oder Hochformat",
    ],
    [
        "Relay / Root UI",
        "Tower + Architect Ende",
        "Interne Archiv- oder Root-Channel-Oberflaeche",
        "Soll sich wie geleakter Systemscreen anfuehlen",
        "PNG/SVG oder Photoshop/Figma-Screen",
    ],
    [
        "Exit Corridor",
        "Redemption Ende",
        "Gang/Notausgang im Morgenblau, zwei Figuren oder Silhouetten",
        "Ein echter Abschluss-Moment, nicht heroisch, eher still",
        "PNG/JPG, 16:9",
    ],
]

SLOT_HEADERS = ["Slot", "Dateinamen", "Hinweis"]
SLOT_ROWS = [
    ["northwing-cctv", "northwing-cctv.(png|jpg|jpeg|mp4|webm)", "Intro / CCTV-Flur"],
    ["desk-dossier-scan", "desk-dossier-scan.(png|jpg|jpeg|mp4|webm)", "Akte / Schreibtisch"],
    ["signal-trace-transcript", "signal-trace-transcript.(png|jpg|jpeg|mp4|webm)", "Transcript / Funkspur"],
    ["relay-ghost-route", "relay-ghost-route.(png|jpg|jpeg|mp4|webm)", "Technik / Relay / optional Video"],
    ["mira-red-room-still", "mira-red-room-still.(png|jpg|jpeg|mp4|webm)", "Mira im roten Raum"],
    ["nullarchive-root-channel", "nullarchive-root-channel.(png|jpg|jpeg|mp4|webm)", "Faux-UI / Root-Channel"],
    ["exit-corridor-dawn", "exit-corridor-dawn.(png|jpg|jpeg|mp4|webm)", "Exit-Flur / Redemption-Ende"],
]

SCRIPT_LINES = [
    ("Intro / CCTV", "02:13. Nordfluegel. Noch ist niemand zu sehen, aber der Fall ist schon im Raum."),
    (
        "Dispatch",
        "Keine Leitung ist fuer diesen Apparat registriert. Wenn das Ding klingelt, ruft niemand Offizielles an.",
    ),
    ("Mira / Rotlicht", "Du suchst mich nicht nur, Elias. Du schreibst mich."),
    ("Architect", "Das System fuehrt Elias Voss nicht nur als Ermittler. Irgendwo steht Autor."),
]

FORMATS = [
    "Bilder: PNG oder JPG, ideal 1920x1080 fuer Querformat-Cutscenes.",
    "Freigestellte Elemente: PNG mit Transparenz.",
    "UI/Archiv-Screens: PNG, SVG oder PSD/Figma-Export.",
    "Audio: WAV oder sauberes MP3, notfalls auch Handyaufnahme als Platzhalter.",
    "Wenn du mir nur Text gibst: bitte jeweils mit Sprecher, Stimmung und ungefaehrer Laenge.",
]

FOOTER = (
    "Kurz gesagt: Wenn du mir 4-6 starke Stills/Scans, 2-4 Text-/Audiofragmente "
    "und eine grobe Look-Richtung gibst, kann ich die neuen Echo-Protocol-Cutscenes "
    "schnell von Placeholder auf glaubwuerdig ziehen."
)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_margins(section) -> None:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    set_margins(section)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(235, 242, 252)

    sub = doc.add_paragraph()
    sub_run = sub.add_run(SUBTITLE)
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(198, 210, 230)

    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = callout.cell(0, 0)
    set_cell_shading(cell, "151A26")
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = para.add_run(PRIORITY_TEXT)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(244, 247, 251)

    doc.add_paragraph("")

    add_docx_heading(doc, "1. Was ich konkret von dir brauche")
    for item in NEEDS:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    add_docx_heading(doc, "2. Shotlist fuer die neue Cutscene-Ebene")
    table = doc.add_table(rows=1, cols=len(SHOT_HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for idx, label in enumerate(SHOT_HEADERS):
        header_cell = table.rows[0].cells[idx]
        header_cell.text = label
        set_cell_shading(header_cell, "1B2435")
        for header_run in header_cell.paragraphs[0].runs:
            header_run.bold = True
            header_run.font.color.rgb = RGBColor(244, 247, 251)

    for row in SHOT_ROWS:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            if idx == 0:
                set_cell_shading(cells[idx], "141A25")

    doc.add_paragraph("")

    add_docx_heading(doc, "3. Mini-Scripts / Textbausteine, die ich fuer Assets und Audio nutzen kann")
    for title_text, body_text in SCRIPT_LINES:
        p = doc.add_paragraph()
        label = p.add_run(f"{title_text}: ")
        label.bold = True
        body = p.add_run(body_text)
        body.italic = True

    doc.add_paragraph("")

    add_docx_heading(doc, "4. Drop-in Slots fuer spaetere Bilder und Videos")
    slot_callout = doc.add_table(rows=1, cols=1)
    slot_callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    slot_cell = slot_callout.cell(0, 0)
    set_cell_shading(slot_cell, "101521")
    slot_para = slot_cell.paragraphs[0]
    slot_run = slot_para.add_run(DROP_FOLDER_TEXT)
    slot_run.font.size = Pt(10.5)
    slot_run.font.color.rgb = RGBColor(244, 247, 251)

    slot_table = doc.add_table(rows=1, cols=len(SLOT_HEADERS))
    slot_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    slot_table.style = "Table Grid"
    for idx, label in enumerate(SLOT_HEADERS):
      header_cell = slot_table.rows[0].cells[idx]
      header_cell.text = label
      set_cell_shading(header_cell, "1B2435")
      for header_run in header_cell.paragraphs[0].runs:
          header_run.bold = True
          header_run.font.color.rgb = RGBColor(244, 247, 251)

    for row in SLOT_ROWS:
      cells = slot_table.add_row().cells
      for idx, value in enumerate(row):
          cells[idx].text = value
          if idx == 0:
              set_cell_shading(cells[idx], "141A25")

    doc.add_paragraph("")

    add_docx_heading(doc, "5. Dateiformate und praktische Lieferung")
    for item in FORMATS:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_run = footer.add_run(FOOTER)
    footer_run.bold = True

    doc.save(DOCX_PATH)


def add_docx_heading(doc: Document, text: str) -> None:
    heading = doc.add_paragraph()
    run = heading.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(126, 216, 255)


def build_pdf() -> None:
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "EchoTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        textColor=colors.HexColor("#ebf2fc"),
        spaceAfter=10,
    )
    body_style = ParagraphStyle(
        "EchoBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        textColor=colors.HexColor("#dbe5f5"),
        spaceAfter=6,
    )
    heading_style = ParagraphStyle(
        "EchoHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#7ed8ff"),
        spaceBefore=10,
        spaceAfter=6,
    )
    label_style = ParagraphStyle(
        "EchoLabel",
        parent=body_style,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#f4f7fb"),
    )
    quote_style = ParagraphStyle(
        "EchoQuote",
        parent=body_style,
        fontName="Helvetica-Oblique",
        textColor=colors.HexColor("#f2d9e2"),
        leftIndent=10,
        borderPadding=6,
    )
    callout_style = ParagraphStyle(
        "EchoCallout",
        parent=body_style,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#f4f7fb"),
        leading=15,
    )

    story = [
        Paragraph(TITLE.replace("\n", "<br/>"), title_style),
        Paragraph(SUBTITLE, body_style),
        Spacer(1, 6),
        Table(
            [[Paragraph(PRIORITY_TEXT.replace("\n", "<br/>"), callout_style)]],
            colWidths=[178 * mm],
            style=TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#151A26")),
                    ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#24324d")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                    ("TOPPADDING", (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                ]
            ),
        ),
        Spacer(1, 10),
        Paragraph("1. Was ich konkret von dir brauche", heading_style),
        ListFlowable(
            [ListItem(Paragraph(item, body_style)) for item in NEEDS],
            bulletType="bullet",
            leftIndent=16,
        ),
        Spacer(1, 8),
        Paragraph("2. Shotlist fuer die neue Cutscene-Ebene", heading_style),
    ]

    shot_table_data = [[Paragraph(header, label_style) for header in SHOT_HEADERS]]
    for row in SHOT_ROWS:
        shot_table_data.append([Paragraph(value, body_style) for value in row])

    shot_table = Table(
        shot_table_data,
        colWidths=[28 * mm, 24 * mm, 47 * mm, 48 * mm, 31 * mm],
        repeatRows=1,
    )
    shot_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1B2435")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#f4f7fb")),
                ("BACKGROUND", (0, 1), (0, -1), colors.HexColor("#141A25")),
                ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#2b3b59")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.extend(
        [
            shot_table,
            Spacer(1, 10),
            Paragraph("3. Mini-Scripts / Textbausteine, die ich fuer Assets und Audio nutzen kann", heading_style),
        ]
    )

    for title_text, body_text in SCRIPT_LINES:
        story.append(Paragraph(f"{title_text}:", label_style))
        story.append(Paragraph(body_text, quote_style))

    story.extend(
        [
            Spacer(1, 10),
            Paragraph("4. Drop-in Slots fuer spaetere Bilder und Videos", heading_style),
            Table(
                [[Paragraph(DROP_FOLDER_TEXT.replace("\n", "<br/>"), callout_style)]],
                colWidths=[178 * mm],
                style=TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#101521")),
                        ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#24324d")),
                        ("LEFTPADDING", (0, 0), (-1, -1), 10),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                        ("TOPPADDING", (0, 0), (-1, -1), 10),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                    ]
                ),
            ),
            Spacer(1, 8),
        ]
    )

    slot_table_data = [[Paragraph(header, label_style) for header in SLOT_HEADERS]]
    for row in SLOT_ROWS:
        slot_table_data.append([Paragraph(value, body_style) for value in row])

    slot_table = Table(
        slot_table_data,
        colWidths=[35 * mm, 92 * mm, 51 * mm],
        repeatRows=1,
    )
    slot_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1B2435")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#f4f7fb")),
                ("BACKGROUND", (0, 1), (0, -1), colors.HexColor("#141A25")),
                ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#2b3b59")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )

    story.extend(
        [
            slot_table,
            Spacer(1, 10),
            Paragraph("5. Dateiformate und praktische Lieferung", heading_style),
            ListFlowable(
                [ListItem(Paragraph(item, body_style)) for item in FORMATS],
                bulletType="bullet",
                leftIndent=16,
            ),
            Spacer(1, 10),
            Paragraph(FOOTER, label_style),
        ]
    )

    doc.build(story)


def main() -> None:
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
